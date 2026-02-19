
from fastapi import FastAPI, File, UploadFile
from fastapi.responses import FileResponse
from fastapi.middleware.cors import CORSMiddleware
from docx import Document
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from pydantic import BaseModel
import os
import shutil
import time

import random
import uuid
from dotenv import load_dotenv

# Load environment variables
load_dotenv()

app = FastAPI()

# Add CORS
app.add_middleware(
    CORSMiddleware,
    allow_origins=["*"],  # Allow all origins for dev
    allow_credentials=True,
    allow_methods=["*"],
    allow_headers=["*"],
)

UPLOAD_DIR = "uploads"
PROCESSED_DIR = "processed"

os.makedirs(UPLOAD_DIR, exist_ok=True)
os.makedirs(PROCESSED_DIR, exist_ok=True)

class ChangeAction(BaseModel):
    action: str 

def create_tracked_insertion(text, author="Agent", date=None):
    if date is None:
        date = time.strftime("%Y-%m-%dT%H:%M:%SZ", time.gmtime())
        
    run = OxmlElement('w:ins')
    # Use UUID masked to positive 30-bit integer for safe ID
    unique_id = str(uuid.uuid4().int & (1<<30)-1)
    run.set(qn('w:id'), unique_id)
    run.set(qn('w:author'), author)
    run.set(qn('w:date'), date)
    
    r = OxmlElement('w:r')
    t = OxmlElement('w:t')
    t.text = text
    r.append(t)
    run.append(r)
    return run

def create_tracked_deletion(text, author="Agent", date=None):
    if date is None:
        date = time.strftime("%Y-%m-%dT%H:%M:%SZ", time.gmtime())
        
    run = OxmlElement('w:del')
    # Use UUID masked to positive 30-bit integer for safe ID
    unique_id = str(uuid.uuid4().int & (1<<30)-1)
    run.set(qn('w:id'), unique_id)
    run.set(qn('w:author'), author)
    run.set(qn('w:date'), date)
    
    r = OxmlElement('w:delText')
    r.text = text
    run.append(r)
    return run


def get_document_changes(doc):
    changes = []
    # Iterate over all paragraphs
    for p in doc.paragraphs:
        # Get all children elements
        children = list(p._p.iterchildren())
        
        i = 0
        while i < len(children):
            child = children[i]
            
            # Check for pattern: Deletion followed by Insertion (Replacement)
            if child.tag.endswith('del'):
                # Look ahead for insertion
                is_replacement = False
                if i + 1 < len(children):
                    next_child = children[i+1]
                    if next_child.tag.endswith('ins'):
                        # Found a pair!
                        del_text = "".join([t.text for t in child.iter() if t.tag.endswith('delText')])
                        ins_text = "".join([t.text for t in next_child.iter() if t.tag.endswith('t')])
                        
                        changes.append({
                            "type": "update",
                            "ids": [child.get(qn('w:id')), next_child.get(qn('w:id'))],
                            "author": child.get(qn('w:author')),
                            "date": child.get(qn('w:date')),
                            "text": f"Change '{del_text}' to '{ins_text}'",
                            "original": del_text,
                            "new": ins_text,
                            "context": p.text[:50] + "..."
                        })
                        i += 2 # Skip both
                        continue

                # If no insertion follows, it's just a deletion
                if not is_replacement:
                    text = "".join([t.text for t in child.iter() if t.tag.endswith('delText')])
                    changes.append({
                        "id": child.get(qn('w:id')),
                        "type": "deletion",
                        "author": child.get(qn('w:author')),
                        "date": child.get(qn('w:date')),
                        "text": text,
                        "context": p.text[:50] + "..."
                    })
                    i += 1
                    
            elif child.tag.endswith('ins'):
                # Standalone insertion (since we handle del+ins above)
                text = "".join([t.text for t in child.iter() if t.tag.endswith('t')])
                changes.append({
                    "id": child.get(qn('w:id')),
                    "type": "insertion",
                    "author": child.get(qn('w:author')),
                    "date": child.get(qn('w:date')),
                    "text": text,
                    "context": p.text[:50] + "..."
                })
                i += 1
            else:
                i += 1
                
    return changes

@app.post("/suggest-changes/")
async def suggest_changes(file: UploadFile = File(...)):
    # Save uploaded file
    file_location = os.path.join(UPLOAD_DIR, file.filename)
    with open(file_location, "wb") as buffer:
        shutil.copyfileobj(file.file, buffer)

    # Process document
    doc = Document(file_location)
    
    # Enable Track Revisions in settings so future user edits are also tracked
    if not doc.settings.element.find(qn('w:trackRevisions')):
        track_revisions = OxmlElement('w:trackRevisions')
        doc.settings.element.append(track_revisions)
    
    # Counter to track modifications
    modifications = 0
    
    for paragraph in doc.paragraphs:
        # Check if it's a header (Heading 1, 2, 3...)
        if paragraph.style.name.startswith('Heading'):
            
            original_text = paragraph.text
            
            # Type 1: Rephrase Suggestion (Replace entire text)
            if modifications % 3 == 0:
                paragraph._p.clear_content()
                del_run = create_tracked_deletion(original_text, author="AI_Reviewer")
                paragraph._p.append(del_run)
                new_text = original_text + " (Refined)"
                ins_run = create_tracked_insertion(new_text, author="AI_Reviewer")
                paragraph._p.append(ins_run)
                
            # Type 2: Punctuation Fix
            elif modifications % 3 == 1:
                paragraph._p.clear_content()
                del_run = create_tracked_deletion(original_text, author="AI_Reviewer")
                paragraph._p.append(del_run)
                new_text = "✨ " + original_text
                ins_run = create_tracked_insertion(new_text, author="AI_Reviewer")
                paragraph._p.append(ins_run)
            
            # Type 3: Uppercase
            elif modifications % 3 == 2:
                paragraph._p.clear_content()
                del_run = create_tracked_deletion(original_text, author="AI_Reviewer")
                paragraph._p.append(del_run)
                new_text = original_text.upper()
                ins_run = create_tracked_insertion(new_text, author="AI_Reviewer")
                paragraph._p.append(ins_run)
            
            modifications += 1

    # Save processed file
    output_filename = f"suggested_{file.filename}"
    output_path = os.path.join(PROCESSED_DIR, output_filename)
    doc.save(output_path)

    return FileResponse(output_path, media_type='application/vnd.openxmlformats-officedocument.wordprocessingml.document', filename=output_filename)

@app.get("/changes/{filename}")
async def list_changes(filename: str):
    file_path = os.path.join(PROCESSED_DIR, filename)
    if not os.path.exists(file_path):
        return {"error": "File not found"}
    
    doc = Document(file_path)
    changes = get_document_changes(doc)
    return {"changes": changes}

@app.post("/changes/{filename}/{change_id}/{action}")
async def handle_change(filename: str, change_id: str, action: str):
    file_path = os.path.join(PROCESSED_DIR, filename)
    if not os.path.exists(file_path):
        return {"error": "File not found"}
    
    doc = Document(file_path)
    
    target_node = None
    
    # helper to find node by ID
    for p in doc.paragraphs:
        for child in p._p.iter():
            if child.get(qn('w:id')) == change_id:
                target_node = child
                break
        if target_node is not None:
            break
            
    if target_node is None:
         return {"error": "Change ID not found"}

    parent = target_node.getparent()
    if parent is None:
        return {"error": "Orphaned node"}

    tag_type = "insertion" if target_node.tag.endswith('ins') else "deletion"

    if action == "accept":
        if tag_type == "insertion":
            # Accept Insertion: Keep content (children of w:ins), remove w:ins wrapper
            index = parent.index(target_node)
            for child in target_node:
                # w:ins -> w:r -> w:t
                # We need to insert these children into the parent
                parent.insert(index, child)
                index += 1
            parent.remove(target_node)
            
        elif tag_type == "deletion":
            # Accept Deletion: Remove the node content effectively (it's already wrapped in del, so we remove the del node)
            parent.remove(target_node)

    elif action == "reject":
        if tag_type == "insertion":
            # Reject Insertion: Remove it entirely
            parent.remove(target_node)
            
        elif tag_type == "deletion":
            # Reject Deletion: Restore content.
            # Extract text from w:delText and create a new run
            text_content = ""
            for node in target_node.iter():
                if node.tag.endswith('delText'):
                    text_content += node.text or ""
            
            new_run = OxmlElement('w:r')
            t = OxmlElement('w:t')
            t.text = text_content
            new_run.append(t)
            
            index = parent.index(target_node)
            parent.insert(index, new_run)
            parent.remove(target_node)


    doc.save(file_path)
    return {"status": "success", "action": action, "id": change_id}


@app.get("/download/{filename}")
async def download_file(filename: str):
    file_path = os.path.join(PROCESSED_DIR, filename)
    if not os.path.exists(file_path):
        return {"error": "File not found"}
    
    # Transform filename for download
    # Internal: suggested_myfile.docx
    # Target: myfile_AI_Suggestions.docx
    
    download_name = filename
    if filename.startswith("suggested_"):
        original_name = filename[10:] # remove "suggested_"
        base, ext = os.path.splitext(original_name)
        download_name = f"{base}_AI_Suggestions{ext}"
    
    return FileResponse(file_path, media_type='application/vnd.openxmlformats-officedocument.wordprocessingml.document', filename=download_name)

if __name__ == "__main__":
    import uvicorn
    uvicorn.run(app, host="0.0.0.0", port=8000)
